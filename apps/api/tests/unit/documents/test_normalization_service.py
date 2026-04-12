from __future__ import annotations

from pathlib import Path

import fitz
from docx import Document as WordDocument
from PIL import Image

from knowledge_chatbox_api.services.documents.normalization_service import NormalizationService


class VisionProviderStub:
    supports_vision = True

    def __init__(self) -> None:
        self.seen_settings = None

    def analyze_image(self, inputs, settings) -> str:
        assert inputs
        self.seen_settings = settings
        return "图片中的文字内容"


class NoVisionProviderStub:
    supports_vision = False

    def analyze_image(self, _inputs, _settings) -> str:
        raise AssertionError("should not be called")


class FailingVisionProviderStub:
    supports_vision = True

    def analyze_image(self, inputs, settings) -> str:
        del inputs, settings
        raise RuntimeError("vision unavailable")


def create_service(tmp_path: Path, provider=None, provider_settings=None) -> NormalizationService:
    return NormalizationService(
        normalized_dir=tmp_path / "normalized",
        provider=provider,
        provider_settings=provider_settings,
    )


def test_txt_is_read_directly(tmp_path: Path) -> None:
    file_path = tmp_path / "sample.txt"
    file_path.write_text("hello\nworld", encoding="utf-8")

    result = create_service(tmp_path).normalize(file_path, "txt")

    assert result.content == "hello\nworld"
    assert result.media_type == "text/plain"
    assert Path(result.normalized_path).read_text(encoding="utf-8") == "hello\nworld"


def test_md_keeps_structure_after_cleanup(tmp_path: Path) -> None:
    file_path = tmp_path / "sample.md"
    file_path.write_text("# Title\n\n- item 1\n- item 2\n", encoding="utf-8")

    result = create_service(tmp_path).normalize(file_path, "md")

    assert "# Title" in result.content
    assert "- item 1" in result.content
    assert result.media_type == "text/markdown"


def test_pdf_extracts_text(tmp_path: Path) -> None:
    file_path = tmp_path / "sample.pdf"
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "PDF content")
    document.save(str(file_path))
    document.close()

    result = create_service(tmp_path).normalize(file_path, "pdf")

    assert "PDF content" in result.content
    assert result.media_type == "text/markdown"


def test_docx_extracts_headings_paragraphs_and_tables(tmp_path: Path) -> None:
    file_path = tmp_path / "sample.docx"
    document = WordDocument()
    document.add_heading("Project Plan", level=1)
    document.add_paragraph("First paragraph")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Name"
    table.rows[0].cells[1].text = "Owner"
    table.rows[1].cells[0].text = "Chatbox"
    table.rows[1].cells[1].text = "Team"
    document.save(str(file_path))

    result = create_service(tmp_path).normalize(file_path, "docx")

    assert "# Project Plan" in result.content
    assert "First paragraph" in result.content
    assert "| Name | Owner |" in result.content
    assert "| Chatbox | Team |" in result.content
    assert result.media_type == "text/markdown"


def test_image_uses_provider_vision_when_supported(tmp_path: Path) -> None:
    file_path = tmp_path / "sample.png"
    Image.new("RGB", (32, 32), color="white").save(file_path)
    provider = VisionProviderStub()
    provider_settings = type("ProviderSettingsStub", (), {"openai_api_key": "dummy-key"})()

    result = create_service(
        tmp_path,
        provider=provider,
        provider_settings=provider_settings,
    ).normalize(file_path, "png")

    assert result.content == "图片中的文字内容"
    assert result.media_type == "text/markdown"
    assert provider.seen_settings is provider_settings


def test_image_falls_back_to_metadata_when_vision_is_unsupported(tmp_path: Path) -> None:
    file_path = tmp_path / "sample.png"
    Image.new("RGB", (32, 32), color="white").save(file_path)

    service = create_service(tmp_path, provider=NoVisionProviderStub())
    result = service.normalize(file_path, "png")

    assert "sample.png" in result.content
    assert "32x32" in result.content
    assert "当前 provider 未提供视觉解析能力" in result.content
    assert result.media_type == "text/markdown"


def test_image_falls_back_to_metadata_when_vision_provider_errors(tmp_path: Path) -> None:
    file_path = tmp_path / "sample.png"
    Image.new("RGB", (64, 32), color="white").save(file_path)

    service = create_service(tmp_path, provider=FailingVisionProviderStub())
    result = service.normalize(file_path, "png")

    assert "sample.png" in result.content
    assert "64x32" in result.content
    assert "vision unavailable" in result.content
