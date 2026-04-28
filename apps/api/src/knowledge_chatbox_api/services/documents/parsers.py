"""文档相关服务模块。"""

from collections.abc import Sequence
from pathlib import Path
from typing import Any, Protocol
from zipfile import BadZipFile

import pymupdf
from docx import Document as WordDocument
from docx.opc.exceptions import PackageNotFoundError
from docx.table import Table
from docx.text.paragraph import Paragraph
from pydantic import BaseModel, ConfigDict

from knowledge_chatbox_api.services.documents.errors import InvalidDocumentError
from knowledge_chatbox_api.utils.image import prepare_image_bytes


class ParsedDocument(BaseModel):
    """文档解析结果。"""

    model_config = ConfigDict(frozen=True)

    content: str
    media_type: str


class DocumentParser(Protocol):
    def parse(self, file_path: Path) -> ParsedDocument: ...


class TextDocumentParser:
    def parse(self, file_path: Path) -> ParsedDocument:
        return ParsedDocument(
            content=file_path.read_text(encoding="utf-8"),
            media_type="text/plain",
        )


class MarkdownDocumentParser:
    def parse(self, file_path: Path) -> ParsedDocument:
        return ParsedDocument(
            content=file_path.read_text(encoding="utf-8").strip(),
            media_type="text/markdown",
        )


class PdfDocumentParser:
    def parse(self, file_path: Path) -> ParsedDocument:
        with pymupdf.open(file_path) as document:
            parts: list[str] = []
            for page in document:
                text: str | list[Any] | dict[str, Any] = page.get_text("text")
                parts.append(text.strip() if isinstance(text, str) else "")

        return ParsedDocument(
            content="\n\n".join(part for part in parts if part).strip(),
            media_type="text/markdown",
        )


class DocxDocumentParser:
    def parse(self, file_path: Path) -> ParsedDocument:
        try:
            document = WordDocument(str(file_path))
        except (BadZipFile, PackageNotFoundError) as exc:
            raise InvalidDocumentError("Invalid or corrupted docx document.") from exc

        blocks: list[str] = []

        for block in document.iter_inner_content():
            if isinstance(block, Paragraph):
                normalized = self._normalize_paragraph(block)
            else:
                normalized = self._normalize_table(block)

            if normalized:
                blocks.append(normalized)

        return ParsedDocument(
            content="\n\n".join(blocks).strip(),
            media_type="text/markdown",
        )

    def _normalize_paragraph(self, paragraph: Paragraph) -> str:
        text = self._collapse_text(paragraph.text)
        if not text:
            return ""

        style = paragraph.style
        style_name = ""
        if style is not None and isinstance(style.name, str):
            style_name = style.name.lower()
        if style_name == "title":
            return f"# {text}"
        if style_name.startswith("heading "):
            suffix = style_name.removeprefix("heading ").strip()
            level = int(suffix) if suffix.isdigit() else 1
            level = max(1, min(level, 6))
            return f"{'#' * level} {text}"
        if "list bullet" in style_name:
            return f"- {text}"
        if "list number" in style_name:
            return f"1. {text}"
        return text

    def _normalize_table(self, table: Table) -> str:
        rows = [self._normalize_row(row.cells) for row in table.rows]
        rows = [row for row in rows if any(cell for cell in row)]
        if not rows:
            return ""

        column_count = max(len(row) for row in rows)
        padded_rows = [row + [""] * (column_count - len(row)) for row in rows]
        header = padded_rows[0]
        divider = ["---"] * column_count
        lines = [self._to_markdown_row(header), self._to_markdown_row(divider)]
        lines.extend(self._to_markdown_row(row) for row in padded_rows[1:])
        return "\n".join(lines)

    def _normalize_row(self, cells: Sequence[Any]) -> list[str]:
        return [self._escape_table_cell(self._collapse_text(cell.text)) for cell in cells]

    def _to_markdown_row(self, cells: list[str]) -> str:
        return f"| {' | '.join(cells)} |"

    def _escape_table_cell(self, value: str) -> str:
        return value.replace("|", "\\|")

    def _collapse_text(self, value: object) -> str:
        if not isinstance(value, str):
            return ""
        return " ".join(value.replace("\xa0", " ").split())


class ImageDocumentParser:
    def __init__(self, *, provider=None, provider_settings=None) -> None:
        self.provider = provider
        self.provider_settings = provider_settings

    def parse(self, file_path: Path) -> ParsedDocument:
        width, height, image_bytes = prepare_image_bytes(file_path, max_dimension=1600)

        if self.provider is None or not getattr(self.provider, "supports_vision", False):
            content = self._build_image_fallback_content(
                file_path=file_path,
                width=width,
                height=height,
                reason="当前 provider 未提供视觉解析能力。",
            )
            return ParsedDocument(content=content, media_type="text/markdown")

        try:
            content = self.provider.analyze_image(
                [{"type": "image", "bytes": image_bytes, "mime_type": "image/jpeg"}],
                self.provider_settings,
            )
        except Exception as exc:
            content = self._build_image_fallback_content(
                file_path=file_path,
                width=width,
                height=height,
                reason=f"视觉解析不可用：{exc}",
            )

        return ParsedDocument(content=content, media_type="text/markdown")

    def _build_image_fallback_content(
        self,
        *,
        file_path: Path,
        width: int,
        height: int,
        reason: str,
    ) -> str:
        return "\n".join(
            [
                f"# 图片资源：{file_path.name}",
                "",
                f"- 文件名：{file_path.name}",
                f"- 格式：{file_path.suffix.lstrip('.').lower()}",
                f"- 分辨率：{width}x{height}",
                "",
                "## 说明",
                reason,
            ]
        ).strip()
